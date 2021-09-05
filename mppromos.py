from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx.shared import Inches
import time
import requests

def get_data(ref_list, document):
        
    path = "C:/chromedriver.exe"
    driver = webdriver.Chrome(path)
    driver.get('chrome://settings/')
    driver.execute_script('chrome.settingsPrivate.setDefaultZoom(0.25);')
    driver.get('https://www.mppromocionales.com/')
    time.sleep(5)

    for reference in ref_list:

        # Busca referencia
        search_input = driver.find_element_by_id('mat-input-0')
        search_input.clear()
        search_input.send_keys(reference)
        search_input.send_keys(Keys.RETURN)
        time.sleep(3)

        # titulo, referencia y descripción
        try:
            title = driver.find_element_by_xpath("//h1[@class='g-font-size-20 g-font-weight-600']")
            title_text = title.text
            titulo = document.add_paragraph()
            titulo.add_run(f'{ref_list.index(reference)+1}.{title_text}').bold = True
        except:
            print(f'No se pudo obtener el título de la ref {reference}')

        try:
            sub_title = driver.find_element_by_xpath("//div[@class='g-font-size-18 g-mb-15']")
            document.add_paragraph(sub_title.text)
        except:
            print(f'No se pudo obtener el subtítulo de la ref {reference}')
        
        try:
            desc = driver.find_elements_by_xpath("//ul[@class='g-mb-16 g-ml-20 g-pl-0 g-font-size-14']/child::li")
            descripcion = document.add_paragraph()
            for element in desc:
                desc_line = descripcion.add_run(element.text)
                desc_line.add_break()

        except:
            print(f'No se pudo obtener la descripción de la ref {reference}')

        # Inventario
        try:
            colores = driver.find_elements_by_xpath("//mat-table[@class='w-100 inventory-tabla mat-table']/child::mat-row")
            q_colores = len(colores)
            tabla_colores = "//mat-table[@class='w-100 inventory-tabla mat-table']"
            inventario = document.add_paragraph()

            for i in range (1, q_colores + 1):
                color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[3]/span[2]").text
                inv_color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[7]/span[2]").text
                inv_line = inventario.add_run(f"{color} - {inv_color}")
                inv_line.add_break()
        except:
            print('No se pudo obtener el inventario')

        # Imagen
        try:
            img = driver.find_element_by_xpath("//img[@class='ng-star-inserted']")
            img_src = img.get_attribute('src')
            response = requests.get(img_src)

            if response.status_code == 200:
                file = open("./images/sample_image.jpg", "wb")
                file.write(response.content)
                file.close()
                document.add_picture("./images/sample_image.jpg", width=Inches(5.25))
                document.add_page_break()
            else:
                print(f'Error al descargar imagen de la ref {reference}, status code({response.status_code})')
        except:
            print(f'Error al descargar imagen de la ref {reference}')       

   
    driver.quit()