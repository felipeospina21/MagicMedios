from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx.shared import Inches
import time
import requests

def get_data(ref_list, document):
        
    path = "C:/chromedriver.exe"
    driver = webdriver.Chrome(path)
    # driver.get('chrome://settings/')
    # driver.execute_script('chrome.settingsPrivate.setDefaultZoom(0.25);')
    driver.get('https://promocionalesnw.com/')
    time.sleep(5)

    try:
        driver.execute_script("document.getElementsByClassName('fancybox-overlay fancybox-overlay-fixed labpopup')[0].style.display = 'none';")
    except:
        print("No se encontro popup")

    for reference in ref_list:

        # Busca referencia
        search_input = driver.find_element_by_id('search_query_top')
        search_input.send_keys(reference)
        search_input.send_keys(Keys.RETURN)
        time.sleep(3)

        # Entra en primer resultado
        result = driver.find_element_by_xpath("//a[@class='product_image']")
        result.click()

        # titulo y descripción
        try:
            header = driver.find_element_by_xpath("//div[@class='pb-center-column  col-xs-12 col-sm-6 col-md-6']")
            header_text = header.text.split("\n")
            title = header_text[0]
            sub_title = header_text[4]
            titulo = document.add_paragraph()
            titulo.add_run(f'{ref_list.index(reference)+1}.{title} {reference}').bold = True
            document.add_paragraph(sub_title)
        except:
            print(f'No se pudo obtener el título de la ref {reference}')
        
        try:
            desc = header_text[5:8]
            descripcion = document.add_paragraph()
            for element in desc:
                desc_line = descripcion.add_run(element)
                desc_line.add_break()

        except:
            print(f'No se pudo obtener la descripción de la ref {reference}')

        # Inventario
        try:
            colores = driver.find_elements_by_xpath("//table[@class='table-bordered']/tbody[1]/child::tr")
            q_colores = len(colores)
            tabla_colores = "//table[@class='table-bordered']/tbody[1]"
            inventario = document.add_paragraph()

            for i in range (1, q_colores + 1):
                color = driver.find_element_by_xpath(f"{tabla_colores}/tr[{i}]/td[1]").text
                inv_color = driver.find_element_by_xpath(f"{tabla_colores}/tr[{i}]/td[5]").text
                inv_line = inventario.add_run(f"{color} - {inv_color}")
                inv_line.add_break()
                print(color, inv_color)
        except:
            print('No se pudo obtener el inventario')

        # Imagen
        try:
            # time.sleep(5)
            img = driver.find_element_by_xpath("//img[@id='bigpic']")
            img_src = img.get_attribute('src')
            response = requests.get(img_src)

            if response.status_code == 200:
                file = open("./images/sample_image.jpg", "wb")
                file.write(response.content)
                file.close()
                document.add_picture("./images/sample_image.jpg", width=Inches(5.25))
                document.add_page_break()
            else:
                print(f'Error al descargar imagen de la ref {reference}, status code({response.status_code})')
        except:
            print(f'Error al descargar imagen de la ref {reference}')       

   
    driver.quit()