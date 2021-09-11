from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx.shared import Inches
import time
import requests

class Get_Data:
# (self,url, search_box_id, first_result_xpath, header_xpath, title_index, sub_title_index,
#     xpath_colores, xpath_tabla_colores, img_xpath,
#     supplier, ref_list, document):

    def __init__(self, url, supplier):
        self.path = "C:/chromedriver.exe"
        self.driver = webdriver.Chrome(self.path)
        self.driver.get('chrome://settings/')
        self.driver.execute_script('chrome.settingsPrivate.setDefaultZoom(0.25);')
        self.driver.get(url)
        time.sleep(5)

    def check_pop_up(self):
        if self.supplier == 'nwpromo':
            try:
                self.driver.execute_script("document.getElementsByClassName('fancybox-overlay fancybox-overlay-fixed labpopup')[0].style.display = 'none';")
            except:
                print("No se encontro popup")

    # for reference in ref_list:

    #     # Busca referencia
    #     search_input = driver.find_element_by_id(search_box_id)
    #     search_input.clear()
    #     search_input.send_keys(reference)
    #     search_input.send_keys(Keys.RETURN)
    #     time.sleep(3)

    #     # Entra en primer resultado
    #     if first_result_xpath != None:
    #         result = driver.find_element_by_xpath(first_result_xpath)
    #         result.click()

    #     # titulo, referencia y descripción
    #     try:
    #         header = driver.find_element_by_xpath(header_xpath)
    #         header_text = header.text.split("\n")
    #         title = header_text[title_index]
    #         sub_title = header_text[sub_title_index]
    #         titulo = document.add_paragraph()
    #         titulo.add_run(f'{ref_list.index(reference)+1}.{title} {reference}').bold = True
    #         document.add_paragraph(sub_title)
    #     except:
    #         print(f'No se pudo obtener el título de la ref {reference}')
        
    #     # Revisar como estandarizar, todos los procedimientos tienen esta parte diferente
    #     try:
    #         desc = driver.find_elements_by_xpath("//ul[@class='g-mb-16 g-ml-20 g-pl-0 g-font-size-14']/child::li")
    #         descripcion = document.add_paragraph()
    #         for element in desc:
    #             desc_line = descripcion.add_run(element.text)
    #             desc_line.add_break()

    #     except:
    #         print(f'No se pudo obtener la descripción de la ref {reference}')

    #     # Inventario
    #     # Validar aplicación para catalogos promo
    #     if supplier != 'promoop':
    #         try:
    #             colores = driver.find_elements_by_xpath(xpath_colores)
    #             q_colores = len(colores)
    #             tabla_colores = xpath_tabla_colores
    #             inventario = document.add_paragraph()

    #             for i in range (1, q_colores + 1):
    #                 color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[3]/span[2]").text
    #                 inv_color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[7]/span[2]").text
    #                 inv_line = inventario.add_run(f"{color} - {inv_color}")
    #                 inv_line.add_break()
    #         except:
    #             print('No se pudo obtener el inventario')

    #     # Imagen
    #     try:
    #         if supplier == 'promoop':
    #             time.sleep(5)

    #         img = driver.find_element_by_xpath(img_xpath)
    #         img_src = img.get_attribute('src')
    #         response = requests.get(img_src)

    #         if response.status_code == 200:
    #             file = open("./images/sample_image.jpg", "wb")
    #             file.write(response.content)
    #             file.close()
    #             document.add_picture("./images/sample_image.jpg", width=Inches(5.25))
    #             document.add_page_break()
    #         else:
    #             print(f'Error al descargar imagen de la ref {reference}, status code({response.status_code})')
    #     except:
    #         print(f'Error al descargar imagen de la ref {reference}')       

   
    # driver.quit()