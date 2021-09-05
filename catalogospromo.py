from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx.shared import Inches
import time
import requests

def get_data(ref_list, document):
        
    path = "C:/chromedriver.exe"
    driver = webdriver.Chrome(path)
    driver.get('https://www.catalogospromocionales.com/')
    time.sleep(3)

    for reference in ref_list:

        # Busca referencia
        search_input = driver.find_element_by_id('productos')
        search_input.clear()
        search_input.send_keys(reference)
        search_btn = driver.find_element_by_xpath("//input[@id='productos']/following-sibling::a")
        search_btn.click()

        # Entra en primer resultado
        result = driver.find_element_by_xpath("//div[@id='backTable']/div[1]/div[1]/a[1]")
        result.click()

        # titulo y descripción
        try:
            header = "//div[@class='hola']"
            title = driver.find_element_by_xpath(f"{header}/h2[1]")
            titulo = document.add_paragraph()
            titulo.add_run(f'{ref_list.index(reference)+1}.{title.text} {reference}').bold = True
        except:
            print(f'No se pudo obtener el titulo de la ref {reference}')

        try:
            desc = driver.find_element_by_xpath(f"{header}/p[2]")
            document.add_paragraph(desc.text)
        except:
            print(f'No se pudo obtener la descripción de la ref {reference}')

        # Cantidad minima
        try:
            table = "//table[@class='table-list']"
            table_col_1 = driver.find_element_by_xpath(f"{table}/tbody[1]/tr[1]/td[1]")
            table_col_2 = driver.find_element_by_xpath(f"{table}/tbody[1]/tr[1]/td[2]")
            col_1_text = table_col_1.text
            col_2_text = table_col_2.text

            paragraph = document.add_paragraph(f'{col_1_text}  ')
            paragraph.add_run(col_2_text).bold = True
        except:
            print(f'No se pudo obtener la cantidad minima de la ref {reference}')

        # Inventario
        try:
            rows = driver.find_elements_by_xpath(f"//tr[@class='titlesRow']/following-sibling::tr")
            rows_len = len(rows)
            table = document.add_table(rows=rows_len, cols=2)
            col1 = table.columns[0]
            col2 = table.columns[1]
            for row in rows:
                rows_list = row.text.split()
                col1.cells[rows.index(row)].text = rows_list[0]
                col2.cells[rows.index(row)].text = rows_list[3]
        except:
            print('No se pudo obtener el inventario')

        # Imagen
        try:
            img = driver.find_element_by_id("img_01")
            img_src = img.get_attribute('src')
            response = requests.get(img_src)
            # print(response.headers)

            if response.status_code == 200:
                file = open("./images/sample_image.jpg", "wb")
                file.write(response.content)
                file.close()
                document.add_picture("./images/sample_image.jpg", width=Inches(5.25))
                document.add_page_break()
            else:
                print(f'Error al descargar imagen de la ref {reference}, status code({response.status_code})')
        except:
            print(f'Error al descargar imagen de la ref {reference}')       

   
    driver.quit()