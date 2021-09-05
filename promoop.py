from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx.shared import Inches
import time
import requests

def get_data(ref_list, document):
        
    path = "C:/chromedriver.exe"
    driver = webdriver.Chrome(path)
    # driver.get('chrome://settings/')
    # driver.execute_script('chrome.settingsPrivate.setDefaultZoom(0.25);')
    driver.get('https://www.promoopcioncolombia.co/')
    time.sleep(5)

    for reference in ref_list:

        # Busca referencia
        search_input = driver.find_element_by_id('q')
        search_input.clear()
        search_input.send_keys(reference)
        search_input.send_keys(Keys.RETURN)
        time.sleep(3)

        # Entra en primer resultado
        result = driver.find_element_by_xpath("//a[@class='img-responsive']")
        result.click()

        # titulo y descripción
        try:
            header = driver.find_element_by_xpath("//td[@class='table-responsive']")
            header_text = header.text.split("\n")
            title = header_text[2]
            sub_title = header_text[3]
            titulo = document.add_paragraph()
            titulo.add_run(f'{ref_list.index(reference)+1}.{title} {reference}').bold = True
            document.add_paragraph(sub_title)
        except:
            print(f'No se pudo obtener el título de la ref {reference}')
        
        try:
            table = "//table[@class='table-hover table-responsive']/tbody[1]"
            desc = driver.find_elements_by_xpath(f"{table}/child::tr")
            descripcion = document.add_paragraph()
            for element in desc:
                desc_line = descripcion.add_run(element.text)
                desc_line.add_break()

        except:
            print(f'No se pudo obtener la descripción de la ref {reference}')

        # # Inventario
        # try:
        #     colores = driver.find_elements_by_xpath("//mat-table[@class='w-100 inventory-tabla mat-table']/child::mat-row")
        #     q_colores = len(colores)
        #     tabla_colores = "//mat-table[@class='w-100 inventory-tabla mat-table']"
        #     inventario = document.add_paragraph()

        #     for i in range (1, q_colores + 1):
        #         color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[3]/span[2]").text
        #         inv_color = driver.find_element_by_xpath(f"{tabla_colores}/mat-row[{i}]/mat-cell[7]/span[2]").text
        #         inv_line = inventario.add_run(f"{color} - {inv_color}")
        #         inv_line.add_break()
        #         print(color, inv_color)
        # except:
        #     print('No se pudo obtener el inventario')

        # Imagen
        try:
            time.sleep(5)
            img = driver.find_element_by_xpath("//div[@id='imgItem']/img")
            img_src = img.get_attribute('src')
            response = requests.get(img_src)

            if response.status_code == 200:
                file = open("./images/sample_image.jpg", "wb")
                file.write(response.content)
                file.close()
                document.add_picture("./images/sample_image.jpg", width=Inches(5.25))
                document.add_page_break()
            else:
                print(f'Error al descargar imagen de la ref {reference}, status code({response.status_code})')
        except:
            print(f'Error al descargar imagen de la ref {reference}')       

   
    driver.quit()